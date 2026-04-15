from __future__ import annotations

import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "documentation"
OUTPUT_DOCX = DOC_DIR / "AI Workfinder dokumentacja aplikacji.docx"
TEST_REPORT_PDF = DOC_DIR / "AI Workfinder raport z testow.pdf"
TMP_IMG_DIR = DOC_DIR / "_docx_assets"

FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


def set_page_margins(section, top=2.0, bottom=1.7, left=2.2, right=2.2):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(9)


def add_table(document: Document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), str(value), bold=(i == 0))
    document.add_paragraph()
    return table


def set_paragraph_border(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "2")
        el.set(qn("w:color"), "CBD5E1")
        pbdr.append(el)
    p_pr.append(pbdr)


def add_code_block(document: Document, text: str):
    p = document.add_paragraph()
    set_paragraph_border(p)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8.5)
    document.add_paragraph()


def add_bullet(document: Document, text: str):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)


def heading(document: Document, text: str, level: int):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return p


def para(document: Document, text: str, center: bool = False):
    p = document.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(10.5)
    return p


def small_para(document: Document, text: str):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    return p


def img_canvas(w=1400, h=850):
    return Image.new("RGB", (w, h), "white")


def fonts():
    return {
        "title": ImageFont.truetype(FONT_BOLD, 28),
        "subtitle": ImageFont.truetype(FONT_BOLD, 20),
        "body": ImageFont.truetype(FONT_REGULAR, 18),
        "small": ImageFont.truetype(FONT_REGULAR, 15),
    }


def draw_multiline(draw, xy, text, font, fill="black", line_gap=4, width_chars=28):
    x, y = xy
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=width_chars) or [""])
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap


def draw_box(draw, x, y, w, h, title, body_lines, fnts, fill="#EFF6FF", border="#1D4ED8"):
    draw.rounded_rectangle((x, y, x + w, y + h), radius=12, outline=border, fill=fill, width=3)
    draw.rounded_rectangle((x, y, x + w, y + 34), radius=12, outline=border, fill="#BFDBFE", width=3)
    draw.text((x + 10, y + 8), title, font=fnts["subtitle"], fill="#0F172A")
    yy = y + 48
    for line in body_lines:
        draw_multiline(draw, (x + 10, yy), line, fnts["small"], width_chars=30)
        yy += 28


def draw_arrow(draw, x1, y1, x2, y2, fill="#334155"):
    draw.line((x1, y1, x2, y2), fill=fill, width=3)
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 > x1:
            draw.line((x2, y2, x2 - 14, y2 - 6), fill=fill, width=3)
            draw.line((x2, y2, x2 - 14, y2 + 6), fill=fill, width=3)
        else:
            draw.line((x2, y2, x2 + 14, y2 - 6), fill=fill, width=3)
            draw.line((x2, y2, x2 + 14, y2 + 6), fill=fill, width=3)
    else:
        if y2 > y1:
            draw.line((x2, y2, x2 - 6, y2 - 14), fill=fill, width=3)
            draw.line((x2, y2, x2 + 6, y2 - 14), fill=fill, width=3)
        else:
            draw.line((x2, y2, x2 - 6, y2 + 14), fill=fill, width=3)
            draw.line((x2, y2, x2 + 6, y2 + 14), fill=fill, width=3)


def draw_actor(draw, x, y, label, fnts):
    draw.ellipse((x - 18, y - 54, x + 18, y - 18), outline="black", width=3)
    draw.line((x, y - 18, x, y + 40), fill="black", width=3)
    draw.line((x - 30, y + 2, x + 30, y + 2), fill="black", width=3)
    draw.line((x, y + 40, x - 28, y + 74), fill="black", width=3)
    draw.line((x, y + 40, x + 28, y + 74), fill="black", width=3)
    draw.text((x - 28, y + 82), label, font=fnts["small"], fill="black")


def make_use_case(path: Path):
    img = img_canvas()
    draw = ImageDraw.Draw(img)
    fnts = fonts()
    draw.rectangle((250, 70, 1150, 710), outline="#64748B", width=3)
    draw.text((600, 20), "Diagram przypadków użycia", font=fnts["title"], fill="#0F172A")
    draw.text((590, 80), "System AI Workfinder", font=fnts["subtitle"], fill="#0F172A")
    draw_actor(draw, 120, 280, "Gość", fnts)
    draw_actor(draw, 1280, 280, "Użytkownik", fnts)
    ellipses = [
        (430, 180, "Rejestracja"),
        (430, 300, "Logowanie"),
        (700, 160, "Edycja profilu"),
        (700, 280, "Upload CV"),
        (700, 400, "Uruchomienie\nwyszukiwania"),
        (980, 160, "Przegląd ofert"),
        (980, 290, "Podgląd aplikacji"),
        (980, 420, "Status i logi\nSSE"),
    ]
    for cx, cy, label in ellipses:
        draw.ellipse((cx - 120, cy - 36, cx + 120, cy + 36), outline="#0F766E", fill="#CCFBF1", width=3)
        draw_multiline(draw, (cx - 78, cy - 12), label, fnts["small"], width_chars=16)
    for tx, ty in [(310, 180), (310, 300)]:
        draw_arrow(draw, 160, 280, tx, ty)
    for tx, ty in [(860, 160), (860, 290), (860, 420), (820, 160), (820, 280), (820, 400)]:
        draw_arrow(draw, 1220, 280, tx, ty)
    img.save(path)


def make_class(path: Path):
    img = img_canvas(h=900)
    draw = ImageDraw.Draw(img)
    fnts = fonts()
    draw.text((480, 20), "Diagram klas i modeli", font=fnts["title"], fill="#0F172A")
    draw_box(draw, 40, 120, 360, 200, "SchematUzytkownik", ["id: str", "email: EmailStr", "name: str | None", "resume_filename: str | None", "gmail_connected: bool"], fnts)
    draw_box(draw, 510, 120, 380, 200, "RepozytoriumUzytkownikow", ["znajdz_po_emailu()", "znajdz_po_id()", "utworz_pusty_profil()", "zaktualizuj()"], fnts, fill="#EEF2FF")
    draw_box(draw, 980, 120, 360, 200, "SearchState", ["status", "jobs_found", "events", "jobs", "summary"], fnts, fill="#F5F3FF")
    draw_box(draw, 40, 470, 360, 220, "Search Routes", ["/api/start_search", "/api/search_status", "/api/jobs", "/api/search_stream"], fnts, fill="#FEF3C7")
    draw_box(draw, 510, 470, 380, 220, "RunInput / RunSummary", ["specialization", "full_name", "resume_path", "apply_results", "llm_trace"], fnts, fill="#ECFCCB")
    draw_box(draw, 980, 470, 360, 220, "JobPosting / ApplyHttpResult", ["url", "title", "raw_snippet", "status_code", "error"], fnts, fill="#FEE2E2")
    draw_arrow(draw, 400, 220, 510, 220)
    draw_arrow(draw, 890, 220, 980, 220)
    draw_arrow(draw, 220, 320, 220, 470)
    draw_arrow(draw, 700, 320, 700, 470)
    draw_arrow(draw, 1160, 320, 1160, 470)
    draw_arrow(draw, 400, 580, 510, 580)
    draw_arrow(draw, 890, 580, 980, 580)
    img.save(path)


def make_data_model(path: Path):
    img = img_canvas(h=860)
    draw = ImageDraw.Draw(img)
    fnts = fonts()
    draw.text((470, 20), "Model danych", font=fnts["title"], fill="#0F172A")
    draw_box(draw, 40, 120, 560, 250, "MongoDB: users", ["_id: ObjectId", "email: str (unique)", "password_hash: str", "name, phone, location", "job_preferences_text", "gmail_connected, resume_filename"], fnts)
    draw_box(draw, 760, 120, 560, 250, "Stan w pamięci: _states[user_id]", ["status, started_at, finished_at", "jobs_found, attempted_apply, applied_ok", "jobs: list[dict]", "summary: dict", "events: list[dict]"], fnts, fill="#FEF3C7")
    draw_box(draw, 40, 470, 560, 180, "Pliki lokalne", ["backend/uploads/<user_id>/<filename>", "config.json z parametrami uruchomienia", "CV referencjonowane przez resume_filename"], fnts, fill="#ECFCCB")
    draw_box(draw, 760, 470, 560, 180, "Dane pośrednie AI", ["jobs -> one_click_jobs -> apply_results", "WorkflowState i RunSummary", "wyniki nie są trwałe w MongoDB"], fnts, fill="#FEE2E2")
    draw_arrow(draw, 600, 245, 760, 245)
    draw_arrow(draw, 320, 370, 320, 470)
    draw_arrow(draw, 1040, 370, 1040, 470)
    img.save(path)


def make_architecture(path: Path):
    img = img_canvas(h=860)
    draw = ImageDraw.Draw(img)
    fnts = fonts()
    draw.text((420, 20), "Architektura rozwiązania", font=fnts["title"], fill="#0F172A")
    draw_box(draw, 40, 120, 220, 140, "Użytkownik", ["przeglądarka", "interakcja z UI"], fnts, fill="#F8FAFC")
    draw_box(draw, 320, 120, 300, 140, "Frontend React", ["App / Routes", "AuthContext / I18n", "centralny klient HTTP"], fnts)
    draw_box(draw, 680, 120, 300, 140, "FastAPI", ["auth / profile", "search / applications", "JWT + CORS"], fnts, fill="#EEF2FF")
    draw_box(draw, 1040, 120, 260, 140, "MongoDB", ["kolekcja users"], fnts, fill="#ECFCCB")
    draw_box(draw, 320, 420, 300, 140, "SSE + search_state", ["status procesu", "event log", "lista wyników"], fnts, fill="#FEF3C7")
    draw_box(draw, 680, 420, 300, 140, "AI Runner", ["LangGraph", "search -> filter -> apply"], fnts, fill="#FEE2E2")
    draw_box(draw, 1040, 420, 260, 140, "Pliki lokalne", ["uploads/", "config.json"], fnts, fill="#F5F3FF")
    draw_box(draw, 680, 650, 620, 120, "Usługi zewnętrzne", ["JustJoin.it", "OpenAI-compatible API", "HTTP apply endpoints"], fnts, fill="#FFF7ED")
    draw_arrow(draw, 260, 190, 320, 190)
    draw_arrow(draw, 620, 190, 680, 190)
    draw_arrow(draw, 980, 190, 1040, 190)
    draw_arrow(draw, 470, 260, 470, 420)
    draw_arrow(draw, 830, 260, 830, 420)
    draw_arrow(draw, 980, 490, 1040, 490)
    draw_arrow(draw, 830, 560, 830, 650)
    draw_arrow(draw, 620, 490, 680, 490)
    img.save(path)


def make_ui(path: Path):
    img = img_canvas(h=960)
    draw = ImageDraw.Draw(img)
    fnts = fonts()
    draw.text((500, 20), "Makiety UI", font=fnts["title"], fill="#0F172A")
    items = [
        (40, 120, "Login", ["Email", "Hasło", "[ Zaloguj ]", "Link: Rejestracja"]),
        (720, 120, "Dashboard", ["specjalizacja / lokalizacja", "limit / max_apply", "status + statystyki", "log stream + lista ofert"]),
        (40, 520, "Profil", ["imię, telefon, lokalizacja", "preferencje pracy", "upload CV", "[ Zapisz ]"]),
        (720, 520, "Applications", ["lista wysłanych aplikacji", "firma, data, status", "link do ogłoszenia"]),
    ]
    for x, y, title, lines in items:
        draw_box(draw, x, y, 600, 280, title, lines, fnts, fill="#F8FAFC")
        draw.rectangle((x + 36, y + 110, x + 560, y + 140), outline="#94A3B8", width=2)
        draw.rectangle((x + 36, y + 155, x + 560, y + 185), outline="#94A3B8", width=2)
        draw.rectangle((x + 36, y + 210, x + 220, y + 250), outline="#1D4ED8", fill="#DBEAFE", width=2)
    img.save(path)


def extract_full_pdf_text(path: Path) -> list[str]:
    reader = PdfReader(str(path))
    return [(page.extract_text() or "").strip() for page in reader.pages]


def build_docx():
    TMP_IMG_DIR.mkdir(exist_ok=True)
    diagram_paths = {
        "use_case": TMP_IMG_DIR / "use_case.png",
        "class": TMP_IMG_DIR / "class.png",
        "data_model": TMP_IMG_DIR / "data_model.png",
        "architecture": TMP_IMG_DIR / "architecture.png",
        "ui": TMP_IMG_DIR / "ui.png",
    }
    make_use_case(diagram_paths["use_case"])
    make_class(diagram_paths["class"])
    make_data_model(diagram_paths["data_model"])
    make_architecture(diagram_paths["architecture"])
    make_ui(diagram_paths["ui"])

    document = Document()
    section = document.sections[0]
    set_page_margins(section)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Workfinder\nDokumentacja aplikacji")
    run.bold = True
    run.font.size = Pt(20)
    document.add_paragraph()
    para(document, "Autorzy: Roman Nadkernychnyi, Maksim Pyanin, Bohdan Hrytsai, Maksym Stepaniuk", center=True)
    para(document, f"Data opracowania: {date.today().strftime('%d.%m.%Y')}", center=True)
    document.add_paragraph()
    para(
        document,
        "Dokument opisuje problem biznesowy, analizę wymagań, projekt rozwiązania, architekturę, "
        "technologie, interfejsy, model danych oraz implementację aplikacji AI Workfinder. "
        "Ostatnią część dokumentu stanowi pełna treść raportu z testów przeniesiona z pliku PDF.",
    )
    document.add_page_break()

    heading(document, "Spis treści", 1)
    for item in [
        "1. Wprowadzenie",
        "2. Opis problemu",
        "3. Projekt i analiza",
        "4. Implementacja",
        "5. Testy",
    ]:
        add_bullet(document, item)
    document.add_page_break()

    heading(document, "1. Wprowadzenie", 1)
    heading(document, "1.1. Cel dokumentu", 2)
    para(document, "Celem dokumentu jest przedstawienie kompletnej dokumentacji aplikacji AI Workfinder: zakresu produktu, założeń analitycznych, wymagań funkcjonalnych i pozafunkcjonalnych, modelu danych, interfejsów, architektury oraz implementacji.")
    heading(document, "1.2. Przyjęte zasady w dokumencie", 2)
    for item in [
        "Priorytet P0 oznacza wymaganie krytyczne, P1 ważne, P2 uzupełniające.",
        "Skróty: UI – interfejs użytkownika, API – interfejs programistyczny, JWT – token uwierzytelniający, LLM – model językowy, SSE – Server-Sent Events.",
        "Opis odnosi się do aktualnej implementacji znajdującej się w repozytorium projektu.",
    ]:
        add_bullet(document, item)
    heading(document, "1.3. Zakres produktu", 2)
    para(document, "AI Workfinder jest webową aplikacją wspierającą użytkownika indywidualnego w wyszukiwaniu pracy. System umożliwia rejestrację i logowanie, uzupełnienie profilu, dodanie CV, skonfigurowanie parametrów wyszukiwania, uruchomienie wieloagentowego pipeline’u AI oraz przegląd ofert i wysłanych aplikacji.")
    heading(document, "1.4. Literatura", 2)
    add_table(document, [
        ["Pozycja", "Zakres wykorzystania"],
        ["Dokumentacja FastAPI", "warstwa API, routingi, zależności, modele wejścia/wyjścia"],
        ["Dokumentacja React i React Router", "warstwa SPA, routing, konteksty i ochrona tras"],
        ["Dokumentacja MongoDB / PyMongo", "trwałe przechowywanie danych użytkowników"],
        ["Dokumentacja LangChain / LangGraph", "orkiestracja wieloagentowego pipeline’u AI"],
        ["Dokumentacja OpenAI-compatible API", "konfiguracja modelu LLM i integracji AI"],
    ])

    heading(document, "2. Opis problemu", 1)
    heading(document, "2.1. Dokładny opis problemu", 2)
    para(document, "Poszukiwanie pracy w branży IT wymaga równoległego przeglądania wielu ofert, oceny ich dopasowania do kompetencji użytkownika oraz ręcznego wypełniania formularzy aplikacyjnych. Proces jest czasochłonny, powtarzalny i podatny na pominięcie atrakcyjnych ofert. AI Workfinder ma ograniczyć ten koszt operacyjny poprzez centralizację profilu użytkownika oraz automatyzację wyszukiwania i części procesu aplikowania.")
    para(document, "Projekt ma charakter proof of concept / prototypu systemu AI-assisted job search. Głównym celem było sprawdzenie możliwości zbudowania pipeline’u złożonego z wyszukiwania ofert, filtrowania 1-click apply oraz automatycznego wysyłania zgłoszeń na podstawie danych użytkownika i przesłanego CV.")
    heading(document, "2.2. Porównanie dostępnych rozwiązań", 2)
    add_table(document, [
        ["Kategoria", "Charakterystyka", "Ograniczenia względem AI Workfinder"],
        ["Ręczne przeglądanie portali pracy", "użytkownik sam filtruje i aplikuje", "wysoki nakład czasu, brak automatyzacji i centralnego statusu"],
        ["Klasyczne agregatory ofert", "prezentują oferty i podstawowe filtry", "zwykle nie automatyzują wysyłki aplikacji"],
        ["Rozszerzenia auto-apply / skrypty", "upraszczają pojedyncze formularze", "zwykle nie mają profilu użytkownika, dashboardu i pełnego pipeline’u"],
        ["AI Workfinder", "łączy profil, AI search, filtrację i monitoring procesu", "ograniczony do jednego źródła ofert i jednej głównej roli użytkownika"],
    ])
    heading(document, "2.3. Możliwości zastosowania praktycznego", 2)
    for item in [
        "Wsparcie kandydata IT w szybkim reagowaniu na nowe oferty pracy.",
        "Automatyzacja rutynowych zadań: wyszukiwania, filtrowania i części procesu aplikowania.",
        "Budowa podstawy pod przyszły system rekomendacyjny lub platformę job-matching.",
        "Walidacja użyteczności integracji LLM + crawler + HTTP apply w jednym produkcie.",
    ]:
        add_bullet(document, item)

    heading(document, "3. Projekt i analiza", 1)
    heading(document, "3.1. Perspektywa produktu", 2)
    para(document, "Produkt jest aplikacją klient-serwer. Frontend w React udostępnia interfejs webowy, backend w FastAPI udostępnia REST API i SSE, a MongoDB przechowuje profile użytkowników. Wyszukiwanie i automatyczne aplikowanie realizowane są przez pipeline AI korzystający z modelu LLM oraz narzędzi HTTP/scrapingowych.")
    heading(document, "3.2. Funkcje produktu", 2)
    for item in [
        "Rejestracja i logowanie użytkownika z użyciem JWT.",
        "Przechowywanie i edycja profilu użytkownika.",
        "Przesyłanie pliku CV do lokalnego magazynu plików backendu.",
        "Konfiguracja procesu wyszukiwania ofert pracy.",
        "Wielostopniowy pipeline AI: search -> filter -> apply.",
        "Strumieniowanie statusu oraz logów procesu wyszukiwania.",
        "Przegląd listy znalezionych ofert i listy wysłanych aplikacji.",
        "Dwujęzyczny interfejs użytkownika (PL/EN).",
    ]:
        add_bullet(document, item)
    heading(document, "3.3. Ograniczenia", 2)
    for item in [
        "Aktualna implementacja koncentruje się na użytkowniku indywidualnym; role firma/admin nie są zaimplementowane w interfejsie ani API.",
        "Źródłem ofert w aktualnej wersji jest portal justjoin.it.",
        "Status wyszukiwania i lista ofert są przechowywane w pamięci procesu, więc restart serwera czyści stan.",
        "Uruchomienie pipeline’u AI wymaga zewnętrznego endpointu kompatybilnego z OpenAI oraz poprawnej konfiguracji .env.",
        "Integracja Gmail jest przygotowana częściowo, ale nie jest obecnie kompletną ścieżką backendową.",
    ]:
        add_bullet(document, item)
    heading(document, "3.4. Aktorzy i charakterystyka użytkowników", 2)
    add_table(document, [
        ["ID", "Nazwa", "Opis"],
        ["GUEST", "Gość", "Osoba niezalogowana. Może założyć konto lub zalogować się do systemu."],
        ["USER", "Użytkownik", "Osoba zalogowana. Zarządza profilem, CV, uruchamia wyszukiwanie i przegląda wyniki."],
        ["AI_PIPELINE", "Pipeline AI", "Aktor techniczny odpowiedzialny za wykonanie procesu search-filter-apply."],
        ["EXTERNAL_SERVICES", "Usługi zewnętrzne", "JustJoin.it oraz endpoint LLM/OpenAI-compatible."],
    ])
    heading(document, "3.5. Obiekty biznesowe", 2)
    add_table(document, [
        ["Obiekt", "Opis"],
        ["Profil użytkownika", "Dane osobowe i preferencje pracy użytkownika oraz referencja do przesłanego CV."],
        ["CV", "Plik PDF/DOC/DOCX przechowywany lokalnie po stronie backendu."],
        ["Oferta pracy", "Opis znalezionego ogłoszenia wraz z URL, źródłem i statusem aplikacji."],
        ["Aplikacja", "Wynik próby wysłania zgłoszenia do oferty."],
        ["Stan wyszukiwania", "Bieżący status procesu AI, liczniki, zdarzenia i lista wyników."],
        ["Konfiguracja uruchomienia", "Parametry biegu AI zapisywane do pliku config.json."],
    ])
    heading(document, "3.6. Wymagania funkcjonalne", 2)
    add_table(document, [
        ["ID", "Wymaganie", "Priorytet"],
        ["FR1", "Rejestracja użytkownika za pomocą adresu e-mail i hasła.", "P0"],
        ["FR2", "Logowanie i uwierzytelnianie żądań za pomocą JWT.", "P0"],
        ["FR3", "Widok i edycja profilu użytkownika.", "P0"],
        ["FR4", "Przesłanie i przechowywanie pliku CV.", "P0"],
        ["FR5", "Uruchomienie procesu wyszukiwania z parametrami.", "P0"],
        ["FR6", "Pipeline AI obejmujący wyszukiwanie, filtrację i apply.", "P0"],
        ["FR7", "Strumień statusu i logów procesu wyszukiwania.", "P1"],
        ["FR8", "Prezentacja listy ofert i listy wysłanych aplikacji.", "P0"],
        ["FR9", "Zmiana języka interfejsu między polskim i angielskim.", "P1"],
        ["FR10", "Centralna warstwa komunikacji HTTP frontend-backend.", "P1"],
    ])
    document.add_picture(str(diagram_paths["use_case"]), width=Inches(6.6))
    small_para(document, "Rysunek 1. Diagram przypadków użycia dla AI Workfinder.")
    heading(document, "3.7. Charakterystyka interfejsów", 2)
    add_table(document, [
        ["ID", "Wymaganie", "Priorytet"],
        ["UI1", "Interfejs działa w przeglądarce jako SPA z trasami publicznymi i chronionymi.", "P0"],
        ["UI2", "Formularze walidują obecność wymaganych danych.", "P0"],
        ["UI3", "UI prezentuje stan procesu wyszukiwania i statystyki.", "P0"],
        ["UI4", "UI aktualizuje log procesu bez pełnego odświeżania strony dzięki SSE.", "P1"],
        ["UI5", "UI oferuje lokalizację PL/EN i spójne nazewnictwo statusów.", "P1"],
    ])
    add_table(document, [
        ["Interfejs", "Opis"],
        ["REST API", "Endpointy /api/register, /api/login, /api/profile, /api/upload_resume, /api/start_search, /api/search_status, /api/jobs, /api/search_stream, /api/applications."],
        ["SSE", "Endpoint /api/search_stream do strumieniowania logów i statusu procesu AI."],
        ["MongoDB", "Połączenie backendu z bazą dokumentową przechowującą kolekcję users."],
        ["OpenAI-compatible API", "Backend AI wykorzystuje ChatOpenAI z parametrami z .env i config.json."],
        ["JustJoin.it / HTTP apply", "Crawler i automatyzacja apply korzystają z zewnętrznych endpointów portalu ofert pracy."],
    ])
    heading(document, "3.8. Wymagania pozafunkcjonalne", 2)
    add_table(document, [
        ["ID", "Nazwa", "Opis", "Priorytet"],
        ["NFR1", "Bezpieczeństwo", "Hasła nie są przechowywane jako tekst jawny; stosowany jest hash PBKDF2 oraz JWT.", "P0"],
        ["NFR2", "Integralność danych", "Adres e-mail użytkownika musi być unikalny w bazie MongoDB.", "P0"],
        ["NFR3", "Użyteczność", "Interfejs powinien być prosty, formularzowy i czytelny.", "P1"],
        ["NFR4", "Obsługa błędów", "System zwraca przewidywalne komunikaty walidacyjne i błędy startu wyszukiwania.", "P0"],
        ["NFR5", "Rozszerzalność", "Architektura umożliwia dodanie kolejnych źródeł ofert i kolejnych agentów.", "P1"],
        ["NFR6", "Przenośność", "Aplikacja działa lokalnie z FastAPI, Vite i MongoDB.", "P1"],
    ])
    heading(document, "3.9. Diagram klas", 2)
    document.add_picture(str(diagram_paths["class"]), width=Inches(6.6))
    small_para(document, "Rysunek 2. Uproszczony diagram klas i modeli systemu.")
    heading(document, "3.10. Model danych i organizacja przechowywania danych", 2)
    para(document, "Projekt nie korzysta z relacyjnej bazy danych. Dane trwałe przechowywane są w dokumentowej bazie MongoDB, natomiast dane operacyjne procesu wyszukiwania są utrzymywane w pamięci aplikacji. CV przechowywane jest w systemie plików backendu, a konfiguracja pojedynczego uruchomienia pipeline’u w pliku config.json.")
    document.add_picture(str(diagram_paths["data_model"]), width=Inches(6.6))
    small_para(document, "Rysunek 3. Model danych trwałych i danych tymczasowych.")
    add_code_block(document, """Przykładowy dokument użytkownika (MongoDB / users):
{
  "_id": "ObjectId(...)",
  "email": "user@example.com",
  "password_hash": "<hash>",
  "name": "Jan Kowalski",
  "phone": "+48 123 456 789",
  "location": "Warszawa",
  "job_preferences_text": "Python backend, remote",
  "gmail_connected": false,
  "resume_filename": "resume.pdf"
}""")
    add_code_block(document, """Przykładowy stan wyszukiwania (w pamięci procesu):
{
  "status": "finished",
  "jobs_found": 20,
  "applied_ok": 3,
  "attempted_apply": 3,
  "events": [{"ts": "...", "level": "INFO", "message": "AI search started"}],
  "jobs": [{ "title": "Python Developer", "apply_url": "...", "application_status": "Applied" }]
}""")
    heading(document, "3.11. Projekt interfejsu użytkownika", 2)
    para(document, "Interfejs został zorganizowany wokół czterech głównych widoków: logowania, panelu wyszukiwania, profilu oraz listy aplikacji. Część zalogowana pracuje w układzie z lewym menu nawigacyjnym.")
    document.add_picture(str(diagram_paths["ui"]), width=Inches(6.6))
    small_para(document, "Rysunek 4. Uproszczone makiety kluczowych ekranów aplikacji.")

    document.add_page_break()
    heading(document, "4. Implementacja", 1)
    heading(document, "4.1. Architektura rozwiązania", 2)
    para(document, "Architektura systemu ma postać klient-serwer. Frontend React odpowiada za routing, formularze, lokalny stan sesji i prezentację danych. Backend FastAPI wystawia API HTTP/SSE, realizuje autoryzację i obsługę plików oraz steruje pipeline’em AI. MongoDB przechowuje wyłącznie dane użytkowników. Wyniki wyszukiwania i logi przebiegu są trzymane w pamięci procesu.")
    document.add_picture(str(diagram_paths["architecture"]), width=Inches(6.6))
    small_para(document, "Rysunek 5. Diagram architektury rozwiązania.")
    heading(document, "4.2. Użyte technologie", 2)
    add_table(document, [
        ["Warstwa", "Technologie"],
        ["Frontend", "React 19, React Router, Vite, axios, CSS, konteksty Auth/I18n"],
        ["Backend API", "Python, FastAPI, Uvicorn, Pydantic"],
        ["Baza danych", "MongoDB, PyMongo"],
        ["Bezpieczeństwo", "JWT (python-jose), Passlib PBKDF2"],
        ["AI / orkiestracja", "langchain-openai, langgraph, ChatOpenAI"],
        ["Scraping / apply", "requests, BeautifulSoup, lxml, własne wrappery HTTP"],
        ["Pliki i konfiguracja", "python-dotenv, lokalny storage backend/uploads, config.json"],
        ["Testy i raportowanie", "pytest, pytest-cov, Vitest, Testing Library, coverage reports"],
    ])
    heading(document, "4.3. Moduły backendowe", 2)
    add_table(document, [
        ["Moduł", "Rola"],
        ["app/api/routes/auth.py", "rejestracja i logowanie użytkownika"],
        ["app/api/routes/profile.py", "pobieranie i aktualizacja profilu, upload CV"],
        ["app/api/routes/search.py", "start procesu AI, status, lista ofert, stream SSE"],
        ["app/api/routes/applications.py", "lista wysłanych aplikacji użytkownika"],
        ["app/repositories/uzytkownik_repo.py", "operacje CRUD na kolekcji users"],
        ["app/services/search_state.py", "stan procesu wyszukiwania w pamięci"],
        ["app/services/ai_runner.py", "walidacja configu, uruchomienie grafu agentów"],
        ["app/services/ai/*", "definicje agentów, grafu i stanów workflow"],
        ["app/tools/*", "crawler, filtr 1-click, wrapper apply HTTP"],
    ])
    heading(document, "4.4. Moduły frontendowe", 2)
    add_table(document, [
        ["Moduł", "Rola"],
        ["src/App.jsx", "definicja tras publicznych i chronionych"],
        ["src/context/AuthContext.jsx", "stan sesji, login, register, logout, pobranie profilu"],
        ["src/context/I18nContext.jsx", "słowniki tłumaczeń PL/EN i zmiana języka"],
        ["src/pages/LoginPage.jsx / RegisterPage.jsx", "formularze wejścia do systemu"],
        ["src/pages/ProfilePage.jsx", "edycja profilu i upload CV"],
        ["src/pages/DashboardPage.jsx", "konfiguracja wyszukiwania, SSE, lista ofert"],
        ["src/pages/ApplicationsPage.jsx", "lista wysłanych aplikacji"],
        ["src/components/MainLayout.jsx", "szkielet części zalogowanej"],
        ["src/api/*.js", "warstwa komunikacji z backendem"],
    ])
    heading(document, "4.5. Integracje zewnętrzne i ograniczenia implementacyjne", 2)
    for item in [
        "Endpoint LLM jest dostarczany zewnętrznie i wskazywany przez OPENAI_BASE_URL oraz OPENAI_API_KEY.",
        "Część apply HTTP zależy od kompatybilności z aktualnym formularzem i endpointami justjoin.it.",
        "Magazyn aplikacji nie jest trwały; po restarcie backendu stan wyszukiwania i lista ofert są tracone.",
        "W kodzie przygotowano zależności pod Celery/Redis i Gmail API, lecz w aktualnej implementacji główna ścieżka działa na FastAPI BackgroundTasks i lokalnym pipeline’ie.",
    ]:
        add_bullet(document, item)

    document.add_page_break()
    heading(document, "5. Testy", 1)
    para(document, "Poniżej zamieszczono pełną treść raportu z testów przeniesioną z pliku „AI Workfinder raport z testow.pdf”.")
    report_pages = extract_full_pdf_text(TEST_REPORT_PDF)
    for idx, page_text in enumerate(report_pages, start=1):
        heading(document, f"5.{idx}. Raport z testów – strona {idx}", 2)
        cleaned = page_text.replace("\x00", "").strip()
        for block in cleaned.split("\n\n"):
            block = " ".join(line.strip() for line in block.splitlines() if line.strip())
            if not block:
                continue
            para(document, block)

    document.save(str(OUTPUT_DOCX))

    for path in diagram_paths.values():
        if path.exists():
            path.unlink()
    if TMP_IMG_DIR.exists() and not any(TMP_IMG_DIR.iterdir()):
        TMP_IMG_DIR.rmdir()

    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
